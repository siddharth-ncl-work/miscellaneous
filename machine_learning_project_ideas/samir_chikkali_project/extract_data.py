import numpy as np
import pandas as pd
import docx
from collections import defaultdict
from tqdm import tqdm

file=docx.Document('Sita_FinalReport_Samir.docx')


"""
print(len(file.paragraphs))
for i,para in enumerate(file.paragraphs[:100]):
  print(i)
  print(para.text)

for i,table in enumerate(file.tables[:4]):
  print(len(table.rows))
  for row in table.rows:
    for cell in row.cells:
      print(cell.text)
  print()

for i,table in enumerate(file.tables[:4]):
  print(len(table.columns))
  for column in table.columns:
    for cell in column.cells:
      print(cell.text)
  print()
"""

def extractTemp(text):
  return text.split('Temp:')[1].split('Press.:')[0].strip()

def extractPress(text):
  return text.split('Press.:')[1].split('Solvent:')[0].strip()

def extractSolvent(text):
  return text.split('Solvent:')[1].split('Time:')[0].strip()

def extractTime(text):
  return text.split('Time:')[1]

def extractConversion(text):
  return text.split('Result:')[1].split('conversion')[0].strip()

def extractEe(text):
  return text.split('conversion')[1].split('ee')[0].strip()

print('#'*40)
print('VERIFICATION:')
total_num_exp=0
for para in file.paragraphs:
  if 'KK-' in para.text:
    total_num_exp+=1
max_num_rows=0
max_num_cols=0
max_num_rows_table_no_list=[]
max_num_cols_table_no_list=[]
num_rows_list=[]
num_cols_list=[]
for table in file.tables:
  num_rows_list.append(len(table.rows))
  num_cols_list.append(len(table.columns))
max_num_rows=np.amax(num_rows_list)
max_num_cols=np.amax(num_cols_list)
max_num_rows_table_no_list=np.argwhere(num_rows_list == max_num_rows).flatten().tolist()
max_num_cols_table_no_list=np.argwhere(num_cols_list == max_num_cols).flatten().tolist()

print('Tables with maximum no. of rows:')
for table_no in max_num_rows_table_no_list:
  print(f'Table No. {table_no}')
  table=file.tables[table_no]
  for row in table.rows:
    for cell in row.cells:
      print(cell.text,end=' | ')
    print()
  print()

print('Experiment Nos. having max rows:')
exp_no_list=[]
for para in file.paragraphs:
  if 'KK-' in para.text:
    exp_no_list.append(para.text)
max_num_rows_exp_nos=[exp_no_list[table_no] for table_no in max_num_rows_table_no_list]
print(max_num_rows_exp_nos)
print()

#calculate total number of reactions and results
num_rxn_cond=0
num_results=0
for para in file.paragraphs:
  if 'KK-' in para.text:
    #print(para.text)
    pass
  elif 'Reaction conditions' in para.text:
    num_rxn_cond+=1
    #print(para.text)
    #print(extractTime(para.text))
  elif 'Result:' in para.text:
    num_results+=1
    #print(para.text)
    #print(extractEe(para.text))

print(f'Table nos. of max rows = {max_num_rows_table_no_list}')
print(f'Table nos. of max columns = {max_num_cols_table_no_list}')
print(f'Maximum no. of Rows = {max_num_rows}')
print(f'Maximum no. of Columns = {max_num_cols}')
print(f'Total no. of Experiments = {total_num_exp}')
print(f'Total no. of Tables = {len(file.tables)}')
print(f'Total no. of Reaction Conditions = {num_rxn_cond}')
print(f'Total no. of Results = {num_results}')
print('#'*40)

data=defaultdict(list)
na='--'

#extract rxn cond
faulty_exp_id_list=['KK-369']
for para in tqdm(file.paragraphs):
  curr_exp_id=None
  if 'KK-' in para.text:
    curr_exp_id=para.text.strip()
    data['exp_id'].append(curr_exp_id)
    if curr_exp_id in faulty_exp_id_list:
      data['Temp'].append(na)
      data['Press'].append(na)
      data['Solvent'].append(na)
      data['Time'].append(na)
  elif 'Reaction conditions' in para.text:
    data['Temp'].append(extractTemp(para.text))
    data['Press'].append(extractPress(para.text))
    data['Solvent'].append(extractSolvent(para.text))
    data['Time'].append(extractTime(para.text))
  elif 'Result:' in para.text:
    data['conversion (%)'].append(extractConversion(para.text))
    data['ee (%)'].append(extractEe(para.text))

#tables
for table in tqdm(file.tables):
  num_rows=len(table.rows)
  for idx,col in enumerate(table.columns):
    if idx==0:
      continue
    key=col.cells[0].text.strip()
    unit=''
    if key == 'µmol':
      key='mols'
      unit='µmol'
    elif key == 'mmol':
      key='mols'
      unit='mmol'
    for i,cell in enumerate(col.cells[1:]):
      data[f'{key}_{i+1}'].append(cell.text.strip()+' '+unit)
    remaining_rows=max_num_rows-num_rows
    for i in range(remaining_rows):
      data[f'{key}_{num_rows+i}'].append(na)


print(f'Columns: {data.keys()}')
df=pd.DataFrame.from_dict(data)
df.to_csv('extracted_data.csv',index=False)
