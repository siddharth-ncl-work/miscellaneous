import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import matplotlib as mpl
import seaborn as sns
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

mpl.rcParams['axes.grid'] = True
mpl.rcParams['figure.figsize']=(15,7)
mpl.rcParams['axes.titlesize']=20
plt.style.use('seaborn-whitegrid')
bar_width=3

raw_data=pd.read_csv('extracted_data.csv')
print(raw_data.head())
print(raw_data.shape)

#remove duplicates
cols=[col for col in raw_data.columns if 'id' not in col]
processed_data=raw_data.drop_duplicates(subset=cols)
print(processed_data.shape)

out_doc = Document()

def plotDist(data,width=bar_width,rotate_x=False):
  plt.figure()
  unique_values=data.value_counts(normalize=True)*100
  plt.bar(unique_values.index,unique_values,width=width,edgecolor='b')
  plt.ylabel('Frequency (%)',fontweight='bold')
  plt.xlabel(data.name,fontweight='bold')
  title='Distribution of '+data.name
  plt.title(title)
  if rotate_x:
    plt.xticks(rotation=90)

  if data.name=='ee (%)':
    plt.title('Distribution of '+data.name+' [Target Variable]')
    plt.axvspan(80, 100, color='y', alpha=0.5, lw=0)
    plt.annotate('~58 % of the data', xy=(81,30), size=15,bbox=dict(boxstyle="round", fc="w"))
  plt.tight_layout()
  plt.savefig('output/'+'_'.join(title.split())+'.jpg')
  if False:
    plt.show()
  out_doc.add_picture('output/'+'_'.join(title.split())+'.jpg',width=Inches(9.5))
  last_paragraph = out_doc.paragraphs[-1] 
  last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  
  plt.close()

#ee
processed_data['ee (%)']=processed_data['ee (%)'].apply(lambda x:str(x).split('%')[0].strip()).astype(np.float)
print((processed_data[processed_data['ee (%)']>=80]).shape[0]*100/processed_data.shape[0])
plotDist(processed_data['ee (%)'],width=1)

#CATEGORICAL
#solvent
plotDist(processed_data['Solvent'],width=0.5,rotate_x=True)

#Comp._1
plotDist(processed_data['Comp._1'],width=0.02,rotate_x=True)

#Comp._2
plotDist(processed_data['Comp._2'],width=0.2,rotate_x=True)

#Comp._3
plotDist(processed_data['Comp._3'],width=0.5,rotate_x=True)

#Comp._4
plotDist(processed_data['Comp._4'],width=0.5,rotate_x=True)

#Comp._5
plotDist(processed_data['Comp._5'],width=0.5,rotate_x=True)


#NUMERICAL
#temp
processed_data['Temp']=processed_data['Temp'].replace('RT','25')
processed_data['Temp']=processed_data['Temp'].replace('--',np.nan)
processed_data['Temp']=processed_data['Temp'].apply(lambda x:str(x).split('°C')[0].strip()).astype(np.float)
processed_data=processed_data.rename(columns={'Temp':'Temp (C)'})
plotDist(processed_data['Temp (C)'])

#pressure
processed_data['Press']=processed_data['Press'].replace('--',np.nan)
processed_data['Press']=processed_data['Press'].apply(lambda x:str(x).split('bar')[0].strip()).astype(np.float)
processed_data=processed_data.rename(columns={'Press':'Press (bar)'})
plotDist(processed_data['Press (bar)'])

#Time
processed_data['Time']=processed_data['Time'].replace('--',np.nan)
processed_data['Time']=processed_data['Time'].apply(lambda x:str(x).split('hrs')[0].strip()).astype(np.float)
processed_data=processed_data.rename(columns={'Time':'Time (hrs)'})
plotDist(processed_data['Time (hrs)'])

#Mol. wt._1
plotDist(processed_data['Mol. wt._1'],width=0.1)

#Mol. wt._2
plotDist(processed_data['Mol. wt._2'],width=10)

#Mol. wt._3
processed_data['Mol. wt._3']=processed_data['Mol. wt._3'].replace('--',np.nan)
processed_data['Mol. wt._3']=processed_data['Mol. wt._3'].astype(np.float)
plotDist(processed_data['Mol. wt._3'],width=10)

#Mol. wt._4
processed_data['Mol. wt._4']=processed_data['Mol. wt._4'].replace('--',np.nan)
processed_data['Mol. wt._4']=processed_data['Mol. wt._4'].astype(np.float)
plotDist(processed_data['Mol. wt._4'],width=5)

#Mol. wt._5
processed_data['Mol. wt._5']=processed_data['Mol. wt._5'].replace('--',np.nan)
processed_data['Mol. wt._5']=processed_data['Mol. wt._5'].astype(np.float)
plotDist(processed_data['Mol. wt._5'],width=5)

def extractMols(text):
  mol=None
  if '--' in text:
    mol=np.nan
  elif 'mmol' in text:
    mol=text.split('mmol')[0]
    if 'x' in mol:
      d=float(mol.split('x')[0])
      e=float(mol.split('x')[1].split('10')[1])
      mol=d*(10**e)
    else:
      mol=float(mol)
  elif 'µmol' in text:
    mol=float(text.split('µmol')[0])*0.001
  if mol>20000:
    mol/=10000
  return mol

#mols_1
processed_data['mols_1']=processed_data['mols_1'].apply(extractMols)
processed_data=processed_data.rename(columns={'mols_1':'mols_1 (mmol)'})
plotDist(processed_data['mols_1 (mmol)'],width=0.5)

#mols_2
processed_data['mols_2']=processed_data['mols_2'].apply(extractMols)
processed_data=processed_data.rename(columns={'mols_2':'mols_2 (mmol)'})
plotDist(processed_data['mols_2 (mmol)'],width=0.1)

#mols_3
processed_data['mols_3']=processed_data['mols_3'].apply(extractMols)
processed_data=processed_data.rename(columns={'mols_3':'mols_3 (mmol)'})
plotDist(processed_data['mols_3 (mmol)'],width=0.005)

#mols_4
processed_data['mols_4']=processed_data['mols_4'].apply(extractMols)
processed_data=processed_data.rename(columns={'mols_4':'mols_4 (mmol)'})
plotDist(processed_data['mols_4 (mmol)'],width=0.1)

#mols_5
processed_data['mols_5']=processed_data['mols_5'].apply(extractMols)
processed_data=processed_data.rename(columns={'mols_5':'mols_5 (mmol)'})
plotDist(processed_data['mols_5 (mmol)'],width=0.005)


def extractEq(text):
  eq=None
  if '--' in text or '9 (to substrate)' in text:
    eq=np.nan
  elif 'x' in text:
    d=float(text.split('x')[0])
    e=float(text.split('x')[1].split('10')[1])
    eq=d*(10**e)
  elif 'eq' in text:
    eq=float(text.split('eq')[0])
  elif 'mol%' in text:
    eq=float(text.split('mol%')[0])*0.001
  else:
    eq=float(text)
  return eq

#Eq. wt._1
processed_data['Eq. wt._1']=processed_data['Eq. wt._1'].apply(extractEq)
plotDist(processed_data['Eq. wt._1'],width=15)

#Eq. wt._2
processed_data['Eq. wt._2']=processed_data['Eq. wt._2'].apply(extractEq)
plotDist(processed_data['Eq. wt._2'],width=0.05)


#Eq. wt._3
processed_data['Eq. wt._3']=processed_data['Eq. wt._3'].apply(extractEq)
plotDist(processed_data['Eq. wt._3'],width=0.08)

#Eq. wt._4
processed_data['Eq. wt._4']=processed_data['Eq. wt._4'].apply(extractEq)
plotDist(processed_data['Eq. wt._4'],width=3)

#Eq. wt._5
processed_data['Eq. wt._5']=processed_data['Eq. wt._5'].apply(extractEq)
plotDist(processed_data['Eq. wt._5'],width=1)


#conversion (%)
processed_data['conversion (%)']=processed_data['conversion (%)'].apply(lambda x:str(x).split('%')[0].strip()).astype(np.float)
plotDist(processed_data['conversion (%)'],width=1)



#ee vs solvent
cols=['Temp (C)','Press (bar)','Time (hrs)','Solvent','Comp._3','Comp._4','Eq. wt._1','Eq. wt._2','Eq. wt._3','Eq. wt._4','Eq. wt._5','conversion (%)']
for col in cols:
  plt.scatter(processed_data[col],processed_data['ee (%)'])
  plt.ylabel('ee (%)',fontweight='bold')
  plt.xlabel(col,fontweight='bold')
  plt.title(f'ee (%) Vs. {col}')
  plt.xticks(rotation=90)
  plt.tight_layout()
  plt.savefig(f"output/ee_vs_{'_'.join(col.split())}.jpg")
  plt.show()
#plt.close()

#correlation plot
import seaborn as sns
sns.pairplot(processed_data[['Temp (C)','Press (bar)','Time (hrs)','Solvent','Comp._3','Comp._4','Eq. wt._1','Eq. wt._2','Eq. wt._3','Eq. wt._4','Eq. wt._5','conversion (%)','ee (%)']])
plt.tight_layout()
plt.savefig('output/correlation_plot.jpg')
plt.show()


#save processed dataframe
processed_data.to_csv('processed_data.csv',index=False)
out_doc.save('Distributions.docx')
